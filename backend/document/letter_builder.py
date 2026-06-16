import os
import re
import subprocess
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

import config
from utils.helpers import sanitize_filename


def _add_para(doc, text, alignment=WD_ALIGN_PARAGRAPH.LEFT, bold=False, space_after=0):
    para = doc.add_paragraph()
    para.alignment = alignment
    para.paragraph_format.line_spacing = 1.15
    para.paragraph_format.space_after = Pt(space_after)
    para.paragraph_format.space_before = Pt(0)
    run = para.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    run.bold = bold
    return para


def _set_margins(doc):
    for section in doc.sections:
        for attr in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
            setattr(section, attr, Inches(1))


def _clean_body(text: str) -> str:
    text = str(text)
    # No em/en dashes in the cover letter — replace with a comma so the prose
    # still reads naturally, then tidy any doubled-up punctuation/space.
    text = re.sub(r" *[—–] *", ", ", text)
    text = re.sub(r"\s+,", ",", text)
    text = re.sub(r",\s*,", ", ", text)
    text = re.sub(r"  +", " ", text)
    text = re.sub(r"\t+", " ", text)
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r" +\n", "\n", text)
    text = re.sub(r"\n +", "\n", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def build_letter_docx(letter_body: str, job: dict, save_pdf: bool = False) -> str:
    """
    Build a cover letter DOCX (and optionally PDF) and return the DOCX path.
    """
    company = job.get("company", "Unknown")
    title = job.get("job_title", "Unknown")
    filename = sanitize_filename(f"{company}_{title}_CoverLetter")
    os.makedirs(config.OUTPUT_LETTER_DIR, exist_ok=True)
    docx_path = os.path.join(config.OUTPUT_LETTER_DIR, f"{filename}.docx")

    doc = Document()
    _set_margins(doc)

    # Header — right aligned
    _add_para(doc, config.USER_NAME, WD_ALIGN_PARAGRAPH.RIGHT)
    _add_para(doc, config.USER_EMAIL, WD_ALIGN_PARAGRAPH.RIGHT)
    _add_para(doc, datetime.now().strftime("%B %d, %Y"), WD_ALIGN_PARAGRAPH.RIGHT)

    # Recipient — left aligned
    _add_para(doc, "Campus Recruitment Team")
    _add_para(doc, company)
    _add_para(doc, "")

    # Subject — centre, bold
    _add_para(doc, f"Application for {company} {title}", WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    _add_para(doc, "")

    # Salutation
    _add_para(doc, f"Dear {company} Recruitment Team,")
    _add_para(doc, "")

    # Body paragraphs
    for para_text in _clean_body(letter_body).split("\n\n"):
        para_text = para_text.strip().replace("\n", " ")
        if para_text:
            p = _add_para(doc, para_text)
            p.paragraph_format.space_after = Pt(10)

    # Sign-off
    _add_para(doc, "")
    _add_para(doc, "Best Regards,")
    _add_para(doc, config.USER_NAME)

    doc.save(docx_path)

    if save_pdf:
        _convert_to_pdf(docx_path)

    return docx_path


def _convert_to_pdf(docx_path: str) -> None:
    pdf_path = docx_path.replace(".docx", ".pdf")
    try:
        from docx2pdf import convert
        subprocess.run(["taskkill", "/f", "/im", "WINWORD.EXE"], capture_output=True)
        convert(docx_path, pdf_path, keep_active=False)
        if os.path.exists(pdf_path):
            print(f"  PDF saved: {pdf_path}")
            return
    except Exception as e:
        print(f"  PDF conversion failed: {e}")
    print("  Tip: open the .docx in Word and export as PDF manually.")
