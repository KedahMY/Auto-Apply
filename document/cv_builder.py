import os
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

import config
from utils.helpers import sanitize_filename


def _add_para(doc, text, alignment=WD_ALIGN_PARAGRAPH.LEFT, bold=False,
              font_size=11, space_after=4):
    para = doc.add_paragraph()
    para.alignment = alignment
    para.paragraph_format.space_after = Pt(space_after)
    para.paragraph_format.space_before = Pt(0)
    run = para.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(font_size)
    run.bold = bold
    return para


def _set_margins(doc):
    for section in doc.sections:
        for attr in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
            setattr(section, attr, Inches(1))


def build_cv_docx(tailored_text: str, job: dict) -> str:
    """
    Build a tailored CV DOCX from plain-text AI output and save it.
    Returns the saved file path.
    """
    company = job.get("company", "Unknown")
    title = job.get("job_title", "Unknown")
    filename = sanitize_filename(f"{company}_{title}_CV")
    os.makedirs(config.OUTPUT_CV_DIR, exist_ok=True)
    out_path = os.path.join(config.OUTPUT_CV_DIR, f"{filename}.docx")

    doc = Document()
    _set_margins(doc)

    # Name header
    _add_para(doc, config.USER_NAME, WD_ALIGN_PARAGRAPH.CENTER, bold=True, font_size=16, space_after=2)
    _add_para(doc, config.USER_EMAIL, WD_ALIGN_PARAGRAPH.CENTER, font_size=11, space_after=10)

    # Parse sections from AI output:
    # A section heading is a line in ALL CAPS (or followed by a colon)
    # Bullet points start with "- "
    for line in tailored_text.splitlines():
        stripped = line.strip()
        if not stripped:
            continue

        is_heading = (
            stripped == stripped.upper()
            and len(stripped) > 2
            and not stripped.startswith("-")
        )

        if is_heading:
            para = _add_para(doc, stripped, bold=True, font_size=12, space_after=2)
            # Underline heading
            para.runs[0].underline = True
        elif stripped.startswith("- "):
            _add_para(doc, stripped, font_size=11, space_after=2)
        else:
            _add_para(doc, stripped, font_size=11, space_after=2)

    doc.save(out_path)
    return out_path
