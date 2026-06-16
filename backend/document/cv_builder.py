import os
import shutil
from docx import Document
from docx.shared import Inches, Pt

import config
from utils.helpers import sanitize_filename

# Progressively shrink to these scale factors until the CV fits one page.
# Finest readable steps first so we keep the largest size that still fits.
_FIT_SCALES = (0.97, 0.94, 0.91, 0.88, 0.85, 0.82, 0.79, 0.76, 0.73, 0.70)


def _set_half_margins(doc) -> None:
    """Half-inch (0.5") margins on all sides — maximises usable space so the
    tailored CV stays on a single page."""
    for section in doc.sections:
        for attr in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
            setattr(section, attr, Inches(0.5))


def _iter_all_paragraphs(doc):
    """Every paragraph in the document body and inside (one level of) tables."""
    yield from doc.paragraphs
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs


def _effective_size(run, para, doc) -> Pt:
    """Resolve a run's font size, falling back through the paragraph style chain
    to the document default. Needed because many runs inherit their size."""
    if run.font.size is not None:
        return run.font.size
    style = para.style
    while style is not None:
        if style.font is not None and style.font.size is not None:
            return style.font.size
        style = getattr(style, "base_style", None)
    normal = doc.styles["Normal"]
    if normal.font.size is not None:
        return normal.font.size
    return Pt(11)


def _scale_doc(doc, scale: float) -> None:
    """Multiply every run's font size and paragraph spacing by `scale`."""
    for para in _iter_all_paragraphs(doc):
        pf = para.paragraph_format
        if pf.space_after is not None:
            pf.space_after = Pt(pf.space_after.pt * scale)
        if pf.space_before is not None:
            pf.space_before = Pt(pf.space_before.pt * scale)
        for run in para.runs:
            run.font.size = Pt(round(_effective_size(run, para, doc).pt * scale, 1))


def _page_count(word, abs_path: str) -> int:
    doc = word.Documents.Open(abs_path, ReadOnly=True)
    try:
        doc.Repaginate()
        return int(doc.ComputeStatistics(2))  # 2 = wdStatisticPages
    finally:
        doc.Close(0)


def _fit_to_one_page(out_path: str) -> None:
    """Best-effort: if the rendered CV exceeds one page, progressively shrink
    fonts/spacing until it fits. Requires Word (pywin32) — silently skipped if
    unavailable so generation still succeeds."""
    try:
        import pythoncom
        import win32com.client
    except ImportError:
        return

    abs_out = os.path.abspath(out_path)
    base = out_path + ".base.docx"
    pythoncom.CoInitialize()
    word = None
    try:
        word = win32com.client.DispatchEx("Word.Application")
        word.Visible = False
        word.DisplayAlerts = False

        if _page_count(word, abs_out) <= 1:
            return

        shutil.copy2(out_path, base)  # unscaled baseline to re-derive each step
        for scale in _FIT_SCALES:
            doc = Document(base)
            _scale_doc(doc, scale)
            doc.save(out_path)
            if _page_count(word, abs_out) <= 1:
                return
        # Still over a page at the smallest scale — keep that smallest version.
    except Exception as exc:
        print(f"  One-page fit skipped: {exc}")
    finally:
        if word is not None:
            word.Quit()
        pythoncom.CoUninitialize()
        if os.path.exists(base):
            os.remove(base)


def _apply_replacement(para, new_text: str) -> None:
    """
    Replace the text content of a paragraph while preserving all paragraph-level
    formatting (style, spacing, indentation, bullets). Only the first run's character
    format is kept; mixed-inline formatting within bullets is not common so this is safe.
    """
    # Strip any accidental leading bullet characters the AI might have included
    clean = new_text.lstrip("•·•–—- ").strip()
    # Restore a leading bullet char if the original paragraph had one in its text
    original = para.text
    if original and original[0] in ("•", "·", "•"):
        clean = original[0] + " " + clean

    if not para.runs:
        para.add_run(clean)
        return
    para.runs[0].text = clean
    for run in para.runs[1:]:
        run.text = ""


def build_cv_docx(replacements: dict, job: dict) -> str:
    """
    Copy the original CV DOCX and surgically apply paragraph-level text replacements.
    All formatting (fonts, styles, layout, columns) from the original is preserved.
    Returns the saved output file path.
    """
    company = sanitize_filename(str(job.get("company", "company")))
    title = sanitize_filename(str(job.get("job_title", "role")))
    filename = f"CV_{company}_{title}.docx"
    os.makedirs(config.OUTPUT_CV_DIR, exist_ok=True)
    out_path = os.path.join(config.OUTPUT_CV_DIR, filename)

    shutil.copy2(config.CV_FILE, out_path)

    doc = Document(out_path)
    _set_half_margins(doc)

    for idx_str, new_text in replacements.items():
        try:
            idx = int(idx_str)
        except (ValueError, TypeError):
            continue
        if new_text and 0 <= idx < len(doc.paragraphs):
            _apply_replacement(doc.paragraphs[idx], new_text)

    doc.save(out_path)
    _fit_to_one_page(out_path)
    return out_path
