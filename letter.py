import pandas as pd
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
from datetime import datetime
import sys
import re
import subprocess
import threading

def clean_text_spacing(text):
    """Clean text spacing issues"""
    if not text:
        return text
    text = str(text)
    text = re.sub(r'  +', ' ', text)  # Multiple spaces to single
    text = re.sub(r'\t+', ' ', text)  # Tabs to spaces
    text = text.replace('\r\n', '\n').replace('\r', '\n')  # Normalize line breaks
    text = re.sub(r' +\n', '\n', text)  # Remove trailing spaces
    text = re.sub(r'\n +', '\n', text)  # Remove leading spaces
    text = re.sub(r'\n{3,}', '\n\n', text)  # Limit paragraph breaks
    return text.strip()

def load_jobs_data():
    """Load and validate jobs data from CSV"""
    try:
        df = pd.read_csv('hkust_job.csv', encoding='utf-8-sig')
        # Create letter column if missing
        if 'letter' not in df.columns:
            df['letter'] = False
            df.to_csv('hkust_job.csv', index=False, encoding='utf-8-sig')
        # Check for coverletter column
        if 'coverletter' not in df.columns:
            print("❌ Error: 'coverletter' column not found in CSV file!")
            sys.exit(1)
        return df
    except FileNotFoundError:
        print("❌ Error: hkust_job.csv file not found!")
        sys.exit(1)
    except Exception as e:
        print(f"❌ Error loading CSV file: {e}")
        sys.exit(1)

def get_ready_jobs(df):
    """Get jobs ready for cover letter generation"""
    jobs_without_letters = df[(df['letter'] == False) | (df['letter'].isna())]
    return jobs_without_letters[
        (jobs_without_letters['coverletter'].notna()) & 
        (jobs_without_letters['coverletter'].astype(str).str.strip() != '')
    ]

def create_formatted_paragraph(doc, text, alignment=WD_ALIGN_PARAGRAPH.LEFT, bold=False):
    """Helper to create formatted paragraphs"""
    para = doc.add_paragraph()
    para.alignment = alignment
    para.paragraph_format.line_spacing = 1.15
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.space_before = Pt(0)
    run = para.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.bold = bold
    return para

def create_cover_letter_document(job_data, cover_letter_content):
    """Create a Word document with proper formatting"""
    doc = Document()
    
    # Set margins
    for section in doc.sections:
        section.top_margin = section.bottom_margin = section.left_margin = section.right_margin = Inches(1)
    
    # Header (right aligned)
    create_formatted_paragraph(doc, "Khoo Thien Zhi", WD_ALIGN_PARAGRAPH.RIGHT)
    create_formatted_paragraph(doc, "tzkhoo@connect.ust.hk", WD_ALIGN_PARAGRAPH.RIGHT)
    create_formatted_paragraph(doc, datetime.now().strftime("%B %d, %Y"), WD_ALIGN_PARAGRAPH.RIGHT)
    
    # Recipient (left aligned)
    create_formatted_paragraph(doc, "Campus Recruitment Team")
    create_formatted_paragraph(doc, job_data.get('company', 'Unknown Company'))
    create_formatted_paragraph(doc, "")  # Spacing
    
    # Subject (center aligned, bold)
    company = job_data.get('company', 'Unknown Company')
    job_title = job_data.get('job_title', 'Unknown Position')
    create_formatted_paragraph(doc, f"Application for {company} {job_title}", WD_ALIGN_PARAGRAPH.CENTER, True)
    create_formatted_paragraph(doc, "")  # Spacing
    
    # Cover letter content
    cleaned_content = clean_text_spacing(cover_letter_content)
    for para_text in cleaned_content.split('\n\n'):
        para_text = para_text.strip()
        if para_text:
            para = create_formatted_paragraph(doc, para_text.replace('\n', ' '))
            para.paragraph_format.space_after = Pt(10)
    
    # Signature
    create_formatted_paragraph(doc, "")  # Spacing
    create_formatted_paragraph(doc, "Best Regards,")
    create_formatted_paragraph(doc, "Khoo Thien Zhi")
    
    return doc

def save_document_word_only(doc, job_data):
    """Save document as Word only (no PDF conversion)"""
    company = job_data.get('company', 'Unknown Company')
    job_title = job_data.get('job_title', 'Unknown Position')
    
    # Create folder and clean filename
    folder_path = "CoverLetter"
    os.makedirs(folder_path, exist_ok=True)
    
    filename = f"{company} {job_title}"
    for char in '<>:"/\\|?*':
        filename = filename.replace(char, '_')
    
    # Save Word document only
    docx_path = os.path.join(folder_path, f"{filename}.docx")
    doc.save(docx_path)
    print(f"✅ Word document saved: {docx_path}")

def save_document_with_pdf(doc, job_data):
    """Save document as Word and PDF with improved error handling"""
    company = job_data.get('company', 'Unknown Company')
    job_title = job_data.get('job_title', 'Unknown Position')
    
    # Create folder and clean filename
    folder_path = "CoverLetter"
    os.makedirs(folder_path, exist_ok=True)
    
    filename = f"{company} {job_title}"
    for char in '<>:"/\\|?*':
        filename = filename.replace(char, '_')
    
    # Save Word document
    docx_path = os.path.join(folder_path, f"{filename}.docx")
    doc.save(docx_path)
    print(f"✅ Word document saved: {docx_path}")
    
    # PDF conversion - try simple method first, then fallback to docx2pdf
    pdf_path = os.path.join(folder_path, f"{filename}.pdf")
    
    # Method 1: Try docx2pdf (most reliable on Windows)
    try:
        from docx2pdf import convert
        print("🔄 Converting to PDF using docx2pdf...")
        
        # Kill existing Word processes
        subprocess.run(["taskkill", "/f", "/im", "WINWORD.EXE"], capture_output=True)
        subprocess.run(["taskkill", "/f", "/im", "winword.exe"], capture_output=True)
        
        # Simple conversion
        convert(docx_path, pdf_path, keep_active=False)
        
        if os.path.exists(pdf_path):
            print(f"✅ PDF document saved: {pdf_path}")
            return
            
    except ImportError:
        print("⚠️  docx2pdf not available, trying pypandoc...")
    except Exception as e:
        print(f"⚠️  docx2pdf failed: {e}, trying pypandoc...")
    
    # Method 2: Try pypandoc without specific engines
    try:
        import pypandoc
        print("🔄 Converting to PDF using pypandoc...")
        
        # Simple conversion without engine specification
        pypandoc.convert_file(docx_path, 'pdf', outputfile=pdf_path)
        
        if os.path.exists(pdf_path):
            print(f"✅ PDF document saved: {pdf_path}")
        else:
            print("⚠️  PDF conversion failed - no output file created")
            
    except ImportError:
        print("⚠️  pypandoc not available")
    except Exception as e:
        print(f"⚠️  pypandoc failed: {e}")
        print("💡 Solutions:")
        print("   1. Manual: Open the .docx file in Word → File → Save As → PDF")
        print("   2. Install docx2pdf: pip install docx2pdf")
        print("   3. Install a PDF engine like wkhtmltopdf and add it to PATH")

def process_jobs(df, num_jobs, save_option):
    """Process selected number of jobs with chosen save option"""
    ready_jobs = get_ready_jobs(df)
    selected_jobs = ready_jobs.head(num_jobs)
    
    save_type = "Word + PDF" if save_option == 1 else "Word only"
    print(f"\n📝 Processing {len(selected_jobs)} job(s) - {save_type} mode...")
    print("-" * 50)
    
    generated_count = 0
    updated_rows = []
    
    for idx, (job_index, job) in enumerate(selected_jobs.iterrows(), 1):
        company = job.get('company', 'Unknown Company')
        job_title = job.get('job_title', 'Unknown Position')
        
        print(f"\n[{idx}/{len(selected_jobs)}] Row {job_index + 2}: {company} - {job_title}")
        
        try:
            # Create document
            doc = create_cover_letter_document(job.to_dict(), str(job.get('coverletter', '')))
            
            # Save based on user choice
            if save_option == 1:
                save_document_with_pdf(doc, job.to_dict())
            else:
                save_document_word_only(doc, job.to_dict())
            
            generated_count += 1
            updated_rows.append(job_index)
            print("✅ Cover letter generated!")
            
        except Exception as e:
            print(f"❌ Error: {e}")
    
    # Update CSV
    if updated_rows:
        df.loc[updated_rows, 'letter'] = True
        df.to_csv('hkust_job.csv', index=False, encoding='utf-8-sig')
        print(f"\n✅ Updated CSV: marked {len(updated_rows)} jobs as completed")
    
    print(f"\n🎉 Generated {generated_count} out of {len(selected_jobs)} cover letters")

def main():
    """Main function"""
    print("🎯 HKUST COVER LETTER GENERATOR")
    print("=" * 50)
    
    # Load data
    df = load_jobs_data()
    ready_jobs = get_ready_jobs(df)
    
    print(f"✅ Loaded {len(df)} jobs from CSV")
    print(f"📊 Status: {df['letter'].sum()} completed, {len(ready_jobs)} ready for generation")
    
    if ready_jobs.empty:
        print("❌ No jobs ready for generation!")
        return
    
    # Show next job
    next_job = ready_jobs.iloc[0]
    print(f"\nNext job (Row {ready_jobs.index[0] + 2}):")
    print(f"Company: {next_job.get('company', 'Unknown')}")
    print(f"Position: {next_job.get('job_title', 'Unknown')}")
    
    # Get number of jobs to process
    while True:
        try:
            num_jobs = int(input(f"\nHow many cover letters to generate? (1-{len(ready_jobs)}): "))
            if 1 <= num_jobs <= len(ready_jobs):
                break
            print(f"Please enter 1-{len(ready_jobs)}")
        except ValueError:
            print("Please enter a valid number")
    
    # Get save option
    print("\nSave Options:")
    print("1. Word + PDF (recommended)")
    print("2. Word only (faster)")
    
    while True:
        try:
            save_choice = int(input("Choose option (1-2): "))
            if save_choice in [1, 2]:
                break
            print("Please enter 1 or 2")
        except ValueError:
            print("Please enter a valid number")
    
    if save_choice == 1:
        print("📄 PDF conversion enabled for all documents")
        save_option = 'both'
    else:
        print("📄 Word documents only (no PDF conversion)")
        save_option = 'word_only'
    
    process_jobs(df, num_jobs, save_option)

if __name__ == "__main__":
    main()